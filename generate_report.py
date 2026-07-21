"""
DOCX report generator - reads content from report_data.json
Pure ASCII driver script to avoid encoding issues on Windows.
"""
import json
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "docs", "xiang_mu_bao_gao.docx")

def build_docx(data):
    doc = Document()

    # Global style
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Microsoft YaHei'
        hs.font.bold = True
        hs.font.color.rgb = RGBColor(0x1A, 0x2E, 0x3D)
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        if level == 1:
            hs.font.size = Pt(18)
        elif level == 2:
            hs.font.size = Pt(14)
        else:
            hs.font.size = Pt(12)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    # --- COVER ---
    for _ in range(6):
        doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(data.get("title", ""))
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x2E, 0x3D)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_p.add_run(data.get("subtitle", ""))
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x55, 0x6B, 0x7A)

    doc.add_paragraph()
    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta_p.add_run(data.get("date", ""))
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # --- TOC ---
    doc.add_heading("目  录", level=1)
    doc.add_paragraph()
    toc_text = data.get("toc", "")
    for line in toc_text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.size = Pt(11)

    doc.add_page_break()

    # --- SECTIONS ---
    def add_para(doc, text, bold=False, indent=True):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.5
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.7)
        r = p.add_run(text)
        r.font.name = 'Microsoft YaHei'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        r.font.size = Pt(11)
        r.bold = bold
        return p

    def add_code(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.left_indent = Cm(1)
        r = p.add_run(text)
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        return p

    def add_bullet(doc, text):
        p = doc.add_paragraph(style='List Bullet')
        p.clear()
        r = p.add_run(text)
        r.font.name = 'Microsoft YaHei'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        r.font.size = Pt(11)
        return p

    def make_table(doc, headers, rows):
        table = doc.add_table(rows=len(rows) + 1, cols=len(headers), style='Light Grid Accent 1')
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = text
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
        for ri, row in enumerate(rows):
            for ci, text in enumerate(row):
                table.rows[ri + 1].cells[ci].text = text
        return table

    for sec in data.get("sections", []):
        doc.add_heading(sec["title"], level=sec.get("level", 1))
        for item in sec.get("items", []):
            t = item["type"]
            if t == "para":
                add_para(doc, item["text"], item.get("bold", False))
            elif t == "bullet":
                add_bullet(doc, item["text"])
            elif t == "code":
                add_code(doc, item["text"])
            elif t == "table":
                make_table(doc, item["headers"], item["rows"])
                doc.add_paragraph("")
            elif t == "heading":
                doc.add_heading(item["text"], level=item.get("level", 2))
            elif t == "page_break":
                doc.add_page_break()

    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    with open("report_data.json", "r", encoding="utf-8") as f:
        data = json.load(f)
    path = build_docx(data)
    print(f"Report generated: {path}")
